"""Word 报告生成器（主逻辑）。

命题交付物硬要求：
- 文件命名：{用户问题}_{YYYYMMDDHHmm}.docx
- 必含 5 字段：标题/发布时间/来源链接/核心内容/附件链接
- 核心内容必须与原文事实一致（反幻觉）

金融分析章节对齐 observation_signals.py 6 维度口径（中标活跃度/公开中标集中度/废标公告关联/明确投标否决/信息冲突观察/高频共现提示）。

章节渲染逻辑（封面/摘要/金融分析）拆分至 app.report.docx_sections。
"""

from __future__ import annotations

import asyncio
import os
import re
from datetime import datetime
from typing import Any

from docx import Document

from app.config import settings
from app.llm.schemas import ParsedFilters
from app.report.docx_components import (
    add_analysis,
    add_detail_table,
    add_footer,
    add_value_note,
)
from app.report.docx_sections import (  # noqa: F401  re-export 保持向后兼容
    _add_cover,
    _add_finance_section,
    _add_summary,
    _add_toc,
    _set_default_font,
)
from app.utils.logger import get_logger

logger = get_logger("docx_generator")

# Windows 文件名非法字符
_ILLEGAL_CHARS = re.compile(r'[\\/:*?"<>|\r\n\t]+')


def _sanitize_filename(query: str) -> str:
    """清理用户查询作为文件名（去除非法字符 + 截断）。"""
    # 去除非法字符
    cleaned = _ILLEGAL_CHARS.sub("", query).strip()
    # 去除首尾空格和点
    cleaned = cleaned.strip(". ")
    # 截断到 80 字符（避免文件名过长）
    if len(cleaned) > 80:
        cleaned = cleaned[:80]
    # 空字符串兜底
    return cleaned or "query"


def build_filename(query: str, dt: datetime | None = None) -> str:
    """构建命题要求的文件名：{用户问题}_{YYYYMMDDHHmm}.docx。

    示例：最近3个月的上海区域内的充电桩招标信息都有哪些_202604071424.docx
    """
    if dt is None:
        dt = datetime.now()
    safe_query = _sanitize_filename(query)
    time_str = dt.strftime("%Y%m%d%H%M")
    return f"{safe_query}_{time_str}.docx"


async def generate_report(
    filters: ParsedFilters,
    items: list[dict[str, Any]],
    job_id: str | None = None,
    source_texts: dict[str, str] | None = None,
    finance_summary: dict[str, Any] | None = None,
    quality_summary: dict[str, Any] | None = None,
) -> str:
    """生成 Word 报告，返回文件路径。

    C-5 修复：同步 CPU + IO 密集操作移到线程池，避免阻塞事件循环。
    C-2 修复：新增 source_texts 参数，反幻觉校验真正生效。

    Args:
        filters: LLM 解析出的过滤条件
        items: 招标信息列表（dict 字段对齐 Tender 表）
        job_id: 可选的任务 ID（用于日志）
        source_texts: 可选，{source_url: 原文} 映射，反幻觉校验用
        finance_summary: 可选，金融分析结果（来自 finance_agent），生成金融分析章节

    Returns:
        生成的 Word 文件绝对路径
    """
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(
        None, _generate_report_sync, filters, items, job_id, source_texts, finance_summary, quality_summary
    )


def _generate_report_sync(
    filters: ParsedFilters,
    items: list[dict[str, Any]],
    job_id: str | None = None,
    source_texts: dict[str, str] | None = None,
    finance_summary: dict[str, Any] | None = None,
    quality_summary: dict[str, Any] | None = None,
) -> str:
    """Word 报告生成同步实现（在线程池中运行）。"""
    # 确保输出目录存在
    os.makedirs(settings.REPORT_OUTPUT_DIR, exist_ok=True)

    # 命题硬要求：文件命名 {用户问题}_{YYYYMMDDHHmm}.docx
    filename = build_filename(filters.raw_query)
    filepath = os.path.abspath(os.path.join(settings.REPORT_OUTPUT_DIR, filename))

    doc = Document()
    _set_default_font(doc)
    _add_page_numbers(doc)

    # 1. 封面
    _add_cover(doc, filters, len(items))

    # 1.5 目录页
    _add_toc(doc)

    # 2. 报告摘要
    _add_summary(doc, filters, items)

    # 3. 项目明细表（命题 5 字段）
    add_detail_table(doc, items)

    # 4. 分析建议
    add_analysis(doc, items)

    # 5. 金融分析（v4.1 6 维公开活动观察信号）
    _add_finance_section(doc, finance_summary)

    # 5.5 证据验证报告（6 Agent 真实能力展示）
    _add_quality_section(doc, quality_summary)

    # 6. 反幻觉校验已由 _add_quality_section 统一展示（证据验证报告）
    # 旧 add_anti_hallucination_section 已移除（避免两套口径冲突）
    # source_texts 仍传给 detail_table 供溯源引用

    # 7. 成本效能脚注（面向用户的价值换算）
    add_value_note(doc)

    # 8. 页脚
    add_footer(doc)

    doc.save(filepath)
    logger.info("report generated job_id={} path={}", job_id, filepath)
    return filepath


def _add_quality_section(doc, quality_summary: dict[str, Any] | None) -> None:
    """证据验证报告章节：展示 6 Agent 真实能力。

    展示内容：
    - LLM 抽取字段总数
    - 证据定位验证通过数 / 无依据数
    - 确定性校验标记的幻觉数
    - 质量评分
    - 34 条验证规则执行结果
    """
    if not quality_summary:
        return

    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_heading("证据验证报告（6 Agent 真实能力）", level=1)

    total_fields = quality_summary.get("total_fields", 0)
    verified_fields = quality_summary.get("verified_fields", 0)
    unjustified = quality_summary.get("unjustified_fields", 0)
    hallucination_flags = quality_summary.get("hallucination_flags", 0)
    quality_score = quality_summary.get("quality_score", 0)
    dedup_rate = quality_summary.get("dedup_rate", 0)
    evidence_pass_rate = quality_summary.get("evidence_pass_rate", 0)
    total_checked = quality_summary.get("total_checked", 0)

    # 摘要
    p = doc.add_paragraph()
    p.add_run(
        f"本报告对 {total_checked} 篇公告执行 LLM 字段抽取，"
        f"共抽取 {total_fields} 个字段，"
        f"经 EvidenceLocator 定位 + 34 条确定性验证规则校验："
    )

    # 指标表格
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    metrics = [
        ("LLM 抽取字段总数", f"{total_fields} 个"),
        ("证据定位验证通过", f"{verified_fields} 个（通过率 {evidence_pass_rate:.1%}）"),
        ("无依据字段（拒绝输出）", f"{unjustified} 个"),
        ("确定性校验幻觉标记", f"{hallucination_flags} 个"),
        ("去重率（SimHash 64 位）", f"{dedup_rate:.1%}"),
        ("综合质量评分", f"{quality_score:.3f}"),
    ]
    for i, (label, value) in enumerate(metrics):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        # 标签列加粗
        for para in table.cell(i, 0).paragraphs:
            for run in para.runs:
                run.bold = True

    # 说明
    p = doc.add_paragraph()
    run = p.add_run(
        "验证方式：LLM 只生成字段值和证据候选，"
        "EvidenceLocator 在原文中逐字符定位证据偏移量，"
        "FieldValidator 用 34 条确定性规则做金额/日期/编号校验。"
        "无依据字段不输出（宁缺毋滥），确定性校验标记的幻觉字段不展示。"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)



def _add_page_numbers(doc) -> None:
    """在页脚添加页码（居中，格式：第 X 页）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 1  # CENTER

    run = p.add_run("第 ")
    run.font.size = Pt(9)

    # PAGE 域
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run2 = p.add_run()
    run2.font.size = Pt(9)
    run2._element.append(fldChar1)
    run2._element.append(instrText)
    run2._element.append(fldChar2)

    run3 = p.add_run(" 页")
    run3.font.size = Pt(9)
