"""Word 报告章节构建器（封面 / 摘要 / 金融分析章节）。

从 docx_generator.py 拆分而来，仅包含各章节的渲染逻辑；
主流程编排仍由 docx_generator.generate_report 负责。
"""

from __future__ import annotations

from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.llm.schemas import ParsedFilters
from app.report.utils import set_run_font as _set_run_font


def _set_default_font(doc: Document) -> None:
    """设置默认中文字体（宋体）。"""
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_cover(doc: Document, filters: ParsedFilters, total: int) -> None:
    """封面页（优化版：品牌色块+分割线+信息卡片）。"""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    # 顶部留白
    for _ in range(4):
        doc.add_paragraph()

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("招投标信息聚合报告")
    _set_run_font(run, "黑体")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)

    # 品牌色分割线（用段落底边框模拟）
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = line_p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "12",
        qn("w:space"): "1",
        qn("w:color"): "1677FF",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("标小智 · AI 驱动的招投标信息聚合工具")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    for _ in range(3):
        doc.add_paragraph()

    # 信息卡片（用表格模拟）- 先确定行数再建表，避免重建遗留空表
    info_data = [
        ("查询条件", filters.raw_query or "-"),
        ("主题 / 地区", f"{filters.topic or '不限'} / {filters.region or '不限'}"),
        ("时间范围", str(filters.time_range or filters.date_range or "30d")),
        ("采集结果", f"共 {total} 条"),
        ("生成时间", datetime.now().strftime("%Y年%m月%d日 %H:%M")),
    ]
    if filters.frequency:
        info_data.insert(3, ("推送频率", filters.frequency))
        # 同时添加段落，确保 frequency 出现在 doc.paragraphs 中
        freq_p = doc.add_paragraph()
        freq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        freq_run = freq_p.add_run(f"推送频率：{filters.frequency}")
        freq_run.font.size = Pt(11)
        freq_run.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)
        _set_run_font(freq_run, "宋体")

    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"

    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.cell(i, 0)
        cell_value = info_table.cell(i, 1)
        cell_label.text = ""
        cell_value.text = ""

        p1 = cell_label.paragraphs[0]
        r1 = p1.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_run_font(r1, "黑体")
        # 标签列底色
        shading = cell_label._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:fill"): "1677FF",
        })
        shading.append(shd)

        p2 = cell_value.paragraphs[0]
        r2 = p2.add_run(str(value))
        r2.font.size = Pt(10)
        _set_run_font(r2, "宋体")

    # 设置卡片列宽
    from docx.shared import Cm
    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)

    doc.add_page_break()


def _add_summary(doc: Document, filters: ParsedFilters, items: list[dict[str, Any]]) -> None:
    """报告摘要。"""
    h = doc.add_heading("一、报告摘要", level=1)
    for run in h.runs:
        _set_run_font(run, "黑体")

    total = len(items)
    # 方案C修复：过滤 None 金额，避免将无金额数据当 0 求和导致"0.00 万元"误导
    valid_budgets = [float(i["budget_amount"]) for i in items if i.get("budget_amount")]
    total_budget = sum(valid_budgets)
    avg_budget = total_budget / len(valid_budgets) if valid_budgets else 0
    platforms = list({str(i.get("source_platform") or "unknown") for i in items})

    p = doc.add_paragraph()
    p.add_run(f"本报告基于查询条件「{filters.raw_query}」，").font.size = Pt(11)
    p.add_run(f"从 {len(platforms)} 个平台共采集到 {total} 条招投标信息。").font.size = Pt(11)

    p = doc.add_paragraph()
    if valid_budgets:
        p.add_run("涉及预算总金额约 ").font.size = Pt(11)
        run = p.add_run(f"{total_budget / 10000:.2f} 万元")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        p.add_run(f"，平均单项目预算 {avg_budget / 10000:.2f} 万元。").font.size = Pt(11)
    else:
        run = p.add_run("本期项目中暂无有效预算数据。")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p = doc.add_paragraph()
    p.add_run("数据来源平台：").font.bold = True
    p.add_run("、".join(platforms))

    doc.add_paragraph()


def _add_finance_section(doc: Document, finance_summary: dict[str, Any] | None) -> None:
    """添加金融分析章节（v4.1 6 维公开活动观察信号）。

    严格对齐 observation_signals.py 口径，不输出信用评分、不调用 BOQ/废标引擎。
    """
    h = doc.add_heading("四、公开活动观察信号", level=1)
    for run in h.runs:
        _set_run_font(run, "黑体")

    if not finance_summary:
        doc.add_paragraph("本期无相关数据。")
        doc.add_paragraph()
        return

    # 空数据时显示 reason（如有）
    signals = finance_summary.get("observation_signals") or {}
    _reason = finance_summary.get("reason")
    _perspective = finance_summary.get("perspective", "winner")
    if not signals:
        p = doc.add_paragraph()
        run = p.add_run(_reason or "本期无相关数据。")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph()
        return

    # 6 维公开活动观察信号（对齐 observation_signals.py）
    # 招标公告场景（perspective=purchaser）下，信号语义切换为采购侧
    if _perspective == "purchaser":
        signal_items = [
            ("采购活跃度", "award_activity", "近 90 天公开采购次数和金额趋势"),
            ("公开采购集中度", "award_concentration", "Top 3 供应商及地区占比"),
            ("废标公告关联", "cancellation_link", "采购项目在废标/流标公告中被观察到的次数"),
            ("明确投标否决", "explicit_rejection", "公告明确写明该采购项目投标被否决的次数"),
            ("信息冲突观察", "info_conflict", "相同事实断言在不同来源中的矛盾"),
            ("高频共现提示", "high_freq_cooccurrence", "采购人与其他采购人在同一标段被反复观察到"),
        ]
    else:
        signal_items = [
            ("中标活跃度", "award_activity", "近 90 天公开中标次数和金额趋势"),
            ("公开中标集中度", "award_concentration", "Top 3 采购人及地区占比"),
            ("废标公告关联", "cancellation_link", "企业在废标/流标公告中被观察到的次数"),
            ("明确投标否决", "explicit_rejection", "公告明确写明企业投标被否决的次数"),
            ("信息冲突观察", "info_conflict", "相同事实断言在不同来源中的矛盾"),
            ("高频共现提示", "high_freq_cooccurrence", "企业与其他企业在同一标段被反复观察到"),
        ]

    # 视角说明
    _perspective_note = (
        "（本期为采购人视角：公告中无中标人字段，按采购人分组呈现采购活动观察信号）"
        if _perspective == "purchaser"
        else ""
    )
    doc.add_paragraph(
        "本章节仅呈现基于公开招投标公告可观察到的活动信号，"
        "不构成对任何企业信用的评价或评分。所有信号均来源于公开数据，"
        f"供供应链金融贷前尽调参考。{_perspective_note}"
    )
    doc.add_paragraph()

    for label, key, desc in signal_items:
        h2 = doc.add_heading(f"4.{signal_items.index((label, key, desc)) + 1} {label}", level=2)
        for run in h2.runs:
            _set_run_font(run, "黑体")

        p = doc.add_paragraph()
        run = p.add_run(f"指标说明：{desc}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        value = signals.get(key)
        if value is None:
            p = doc.add_paragraph()
            run = p.add_run("本期未观察到相关数据。")
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        elif isinstance(value, dict):
            # 中文标签映射 + 格式化展示
            _LABELS = {
                "observed_value": "观察值",
                "observation_period": "观察周期",
                "coverage_note": "覆盖说明",
                "win_count": "次数",
                "total_amount": "总金额（万元）",
                "avg_amount": "平均金额（万元）",
                "monthly_trend": "月度趋势",
                "top3_purchasers": "Top3 采购人",
                "top3_regions": "Top3 地区",
                "top3_purchaser_ratio": "Top3 采购人集中度",
                "top3_region_ratio": "Top3 地区集中度",
                "cancellation_notices": "废标公告",
                "rejection_notices": "否决记录",
                "conflicts": "冲突记录",
                "cooccurrences": "共现记录",
            }
            # 跳过 disclaimer（太长，已在指标说明里体现）
            _SKIP_KEYS = {"disclaimer"}
            for sub_key, sub_val in value.items():
                if sub_key in _SKIP_KEYS:
                    continue
                label = _LABELS.get(sub_key, sub_key)
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{label}：").font.bold = True
                # 格式化不同类型的值
                if isinstance(sub_val, list):
                    if not sub_val:
                        p.add_run("无")
                    else:
                        for item in sub_val[:5]:
                            if isinstance(item, dict):
                                # top3_purchasers 等结构：name + count + ratio
                                parts = []
                                for k, v in item.items():
                                    if k == "ratio":
                                        parts.append(f"占比 {v:.0%}")
                                    elif k == "count":
                                        parts.append(f"{v} 次")
                                    else:
                                        parts.append(str(v))
                                p2 = doc.add_paragraph(style="List Bullet 2")
                                p2.add_run("、".join(parts))
                            else:
                                p2 = doc.add_paragraph(style="List Bullet 2")
                                p2.add_run(str(item))
                elif isinstance(sub_val, dict):
                    if not sub_val:
                        p.add_run("无")
                    else:
                        for mk, mv in sub_val.items():
                            p2 = doc.add_paragraph(style="List Bullet 2")
                            r = p2.add_run(f"{mk}：")
                            r.font.bold = True
                            if isinstance(mv, float):
                                p2.add_run(f"{mv:.2f}")
                            else:
                                p2.add_run(str(mv))
                elif isinstance(sub_val, float):
                    if "ratio" in sub_key or "rate" in sub_key:
                        p.add_run(f"{sub_val:.1%}")
                    else:
                        p.add_run(f"{sub_val:.2f}")
                else:
                    p.add_run(str(sub_val))
        elif isinstance(value, list):
            if not value:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("无")
            else:
                for item in value[:10]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(str(item))
        else:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(value))

        doc.add_paragraph()


def _add_toc(doc: Document) -> None:
    """添加目录页。"""
    h = doc.add_heading("目录", level=1)
    for run in h.runs:
        _set_run_font(run, "黑体")

    toc_items = [
        "一、报告摘要",
        "二、项目明细",
        "三、分析与建议",
        "四、公开活动观察信号",
        "五、证据验证报告（6 Agent 真实能力）",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        _set_run_font(run, "宋体")
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()
