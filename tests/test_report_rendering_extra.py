"""渲染层金额处理测试（方案C 修复 Bug 3/4/5/14）。

覆盖目标：
- docx_sections._add_summary: 全 None 金额不显示"0.00 万元"
- docx_components.add_analysis: 全 None 金额跳过高预算建议
- docx_components.add_analysis: 混合金额时 top3 来自有效金额项
"""

from __future__ import annotations

from docx import Document

from app.llm.schemas import ParsedFilters
from app.report.docx_components import add_analysis
from app.report.docx_sections import _add_summary


def _make_filters(query: str = "医疗设备采购") -> ParsedFilters:
    return ParsedFilters(raw_query=query, topic=query)


class TestAddSummaryBudget:
    """方案C Bug 3/14：摘要页金额 None→0 误导修复。"""

    def test_all_none_budget_shows_no_budget_message(self) -> None:
        """全 None 金额 → 显示"暂无有效预算数据"，不显示"0.00 万元"。"""
        doc = Document()
        items = [
            {"project_name": "项目A", "budget_amount": None, "source_platform": "ccgp"},
            {"project_name": "项目B", "budget_amount": None, "source_platform": "ccgp"},
        ]
        _add_summary(doc, _make_filters(), items)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "约 0.00 万元" not in text, "不应显示'约 0.00 万元'（None 不应当 0 求和）"
        assert "暂无有效预算数据" in text, "应显示暂无有效预算数据"

    def test_partial_none_budget_only_sums_valid(self) -> None:
        """混合金额（部分 None）→ 只对有效金额求和，不报 0.00。"""
        doc = Document()
        items = [
            {"project_name": "项目A", "budget_amount": 100000, "source_platform": "ccgp"},
            {"project_name": "项目B", "budget_amount": None, "source_platform": "ccgp"},
            {"project_name": "项目C", "budget_amount": 200000, "source_platform": "ccgp"},
        ]
        _add_summary(doc, _make_filters(), items)
        text = "\n".join(p.text for p in doc.paragraphs)

        # 总预算 = 300000 / 10000 = 30.00 万元
        assert "30.00 万元" in text, "应只对有效金额求和（10万+20万=30万）"
        # 平均 = 300000 / 2 = 150000 / 10000 = 15.00 万元（按有效项数算平均）
        assert "15.00 万元" in text, "平均应按有效金额项数计算"
        # 不应出现"暂无有效预算数据"（因为有有效金额）
        assert "暂无有效预算数据" not in text

    def test_all_valid_budget_normal_display(self) -> None:
        """全有效金额 → 正常显示总额和平均。"""
        doc = Document()
        items = [
            {"project_name": "项目A", "budget_amount": 500000, "source_platform": "ccgp"},
            {"project_name": "项目B", "budget_amount": 300000, "source_platform": "ccgp"},
        ]
        _add_summary(doc, _make_filters(), items)
        text = "\n".join(p.text for p in doc.paragraphs)

        # 总 = 800000 / 10000 = 80.00 万元
        assert "80.00 万元" in text
        # 平均 = 400000 / 10000 = 40.00 万元
        assert "40.00 万元" in text


class TestAddAnalysisBudget:
    """方案C Bug 4/5：高预算建议选 0 金额项目修复。"""

    def test_all_none_budget_skips_high_budget_suggestion(self) -> None:
        """全 None 金额 → 不输出"高预算项目关注建议"。"""
        doc = Document()
        items = [
            {"project_name": "项目A", "budget_amount": None, "tender_org": "甲方A"},
            {"project_name": "项目B", "budget_amount": None, "tender_org": "甲方B"},
        ]
        add_analysis(doc, items)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "高预算项目关注建议" not in text, "无有效金额时不应输出高预算建议"
        # 仍应有数据声明
        assert "数据声明" in text

    def test_partial_none_budget_top3_from_valid(self) -> None:
        """混合金额 → top3 仅来自有效金额项，None 项不入选。"""
        doc = Document()
        items = [
            {"project_name": "有效A", "budget_amount": 100000, "tender_org": "甲方A"},
            {"project_name": "无效B", "budget_amount": None, "tender_org": "甲方B"},
            {"project_name": "有效C", "budget_amount": 500000, "tender_org": "甲方C"},
            {"project_name": "无效D", "budget_amount": None, "tender_org": "甲方D"},
        ]
        add_analysis(doc, items)
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "高预算项目关注建议" in text
        # 最高预算是 有效C（50万），应排第一
        assert "有效C" in text
        assert "有效A" in text
        # 简化断言：有效C 排在有效A 前面（按金额降序）
        idx_c = text.find("有效C")
        idx_a = text.find("有效A")
        assert 0 <= idx_c < idx_a, "有效C（50万）应排在有效A（10万）前面"

    def test_suggestion_numbering_dynamic(self) -> None:
        """建议编号动态递增：无金额时"即将截止"应为 1，"数据声明"应为 2。"""
        doc = Document()
        items = [
            {
                "project_name": "项目",
                "budget_amount": None,
                "deadline": "2026-12-31T23:59",
            },
        ]
        add_analysis(doc, items)
        text = "\n".join(p.text for p in doc.paragraphs)

        # 无金额 → 跳过高预算建议
        # 有 deadline → "1. 即将截止项目提醒"
        assert "1. 即将截止项目提醒" in text, "无金额时即将截止应为第1项"
        # 数据声明应为第2项
        assert "2. 数据声明" in text, "数据声明编号应递增为2"
