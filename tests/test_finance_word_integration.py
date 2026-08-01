"""finance_agent → Word 报告链路集成测试。

验证 _add_finance_section 在 None / 空字典 / 完整数据下的渲染行为，
以及 delivery_agent 是否将 state.finance_summary 透传给 generate_report。
"""
from __future__ import annotations

from types import SimpleNamespace

from docx import Document

from app.report.docx_generator import _add_finance_section


def _all_text(doc: Document) -> str:
    """汇总文档所有段落与表格文本，便于断言。"""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_add_finance_section_with_none():
    """finance_summary=None 时应输出"本期无相关数据"且不抛异常。"""
    doc = Document()
    _add_finance_section(doc, None)  # type: ignore[arg-type]

    text = _all_text(doc)
    assert "本期无相关数据" in text
    assert "金融分析" in text


def test_add_finance_section_with_empty_dict():
    """finance_summary={}（falsy）时同样走空数据分支，不抛异常。"""
    doc = Document()
    _add_finance_section(doc, {})

    text = _all_text(doc)
    assert "本期无相关数据" in text


def test_add_finance_section_with_full_data():
    """完整 finance_summary 应渲染金融分析章节及 BOQ/风险/供应商子标题。"""
    finance_summary = {
        "boq_anomalies": 2,
        "risk_items": 2,
        "avg_supplier_score": 82.0,
        "boq_report": {
            "reports": [{"score": 75.0}, {"score": 80.0}],
            "items": [
                {
                    "name": "服务器",
                    "quantity": 2,
                    "unit": "台",
                    "unit_price": 200000.0,
                    "status": "overpriced",
                },
                {
                    "name": "电脑",
                    "quantity": 5,
                    "unit": "台",
                    "unit_price": 100.0,
                    "status": "underpriced",
                },
            ],
        },
        "risk_report": {
            "reports": [{"risk_score": 45.0}],
            "items": [
                {
                    "clause": "必须具备 AAA 资质",
                    "risk_level": "medium",
                    "law_ref": "《招标投标法》第二十条",
                    "suggestion": "核查排他性条款",
                },
                {
                    "clause": "履约保证金 30%",
                    "risk_level": "high",
                    "law_ref": "",
                    "suggestion": "建议降至 10% 以内",
                },
            ],
        },
        "supplier_scores": [
            {
                "normalized_name": "测试供应商A",
                "dimensions": [
                    {"name": "concentration", "score": 80.0},
                    {"name": "amount_anomaly", "score": 75.0},
                    {"name": "frequency", "score": 90.0},
                    {"name": "region", "score": 85.0},
                    {"name": "purchaser", "score": 70.0},
                ],
                "total_score": 82.0,
                "risk_level": "medium",
            }
        ],
    }

    doc = Document()
    _add_finance_section(doc, finance_summary)

    text = _all_text(doc)
    assert "金融分析" in text
    assert "BOQ 报价异常检测" in text
    assert "废标风险预警" in text
    assert "供应商风险评分" in text
    assert "测试供应商A" in text


async def test_delivery_agent_passes_finance_summary(monkeypatch):
    """delivery_agent 应从 state 取出 finance_summary 并透传给 generate_report。"""
    from app.agents import delivery_agent as da
    from app.scheduler import subscription as sub_mod

    finance_summary = {
        "boq_anomalies": 1,
        "risk_items": 2,
        "avg_supplier_score": 80.0,
    }

    fake_tender = SimpleNamespace(
        project_name="测试项目",
        publish_time=None,
        source_url="http://example.com/t/1",
        core_content="项目正文",
        attachment_url=None,
        budget_amount=None,
        tender_org="某采购单位",
        deadline=None,
        source_platform="ccgp",
    )

    async def _fake_get_unpushed(db, sub_id, parsed):
        return [fake_tender]

    captured: dict = {}

    async def _fake_generate_report(filters, items, **kwargs):
        captured["finance_summary"] = kwargs.get("finance_summary")
        captured["job_id"] = kwargs.get("job_id")
        return "/tmp/fake_report.docx"

    async def _fake_trigger_push(state, sub_id, unpushed):
        return {
            "delivered": False,
            "email_sent": False,
            "webhook_sent": False,
            "message_id": None,
        }

    monkeypatch.setattr(sub_mod, "get_unpushed_tenders", _fake_get_unpushed)
    monkeypatch.setattr(
        "app.report.docx_generator.generate_report", _fake_generate_report
    )
    monkeypatch.setattr(da, "_trigger_push", _fake_trigger_push)

    state = {
        "parsed_filters": SimpleNamespace(),
        "subscription_id": 999,
        "finance_summary": finance_summary,
    }
    result = await da.delivery_agent(state)

    # generate_report 被调用，且 finance_summary 与 state 中的是同一对象
    assert "finance_summary" in captured
    assert captured["finance_summary"] is finance_summary
    # delivery_agent 把报告路径写回 state
    assert result["report_path"] == "/tmp/fake_report.docx"
