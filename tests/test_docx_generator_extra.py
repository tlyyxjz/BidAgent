"""Extra tests for app/report/docx_generator.py.

Covers _add_finance_section branches (lines 119-123, 185-236):
- finance_summary = None
- finance_summary with no observation_signals
- signal value is dict / list / scalar / None
- cover with frequency field
- build_filename / _sanitize_filename edge cases
"""

from __future__ import annotations

import os
from datetime import datetime
from unittest.mock import patch

import pytest
from docx import Document

from app.llm.schemas import ParsedFilters
from app.report.docx_generator import (
    _add_cover,
    _add_finance_section,
    _sanitize_filename,
    build_filename,
    _generate_report_sync,
)


class TestAddFinanceSectionEmpty:
    def test_finance_summary_none(self) -> None:
        doc = Document()
        _add_finance_section(doc, None)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "本期无相关数据" in text

    def test_finance_summary_empty_dict(self) -> None:
        doc = Document()
        _add_finance_section(doc, {})
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "本期无相关数据" in text

    def test_finance_summary_no_observation_signals(self) -> None:
        doc = Document()
        _add_finance_section(doc, {"other_key": "val"})
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "本期无相关数据" in text

    def test_finance_summary_empty_signals(self) -> None:
        doc = Document()
        _add_finance_section(doc, {"observation_signals": {}})
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "本期无相关数据" in text


class TestAddFinanceSectionSignals:
    def test_signal_dict_value(self) -> None:
        doc = Document()
        _add_finance_section(doc, {
            "observation_signals": {
                "award_activity": {"次数": 5, "金额": "100万元"},
            }
        })
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "次数" in text
        assert "5" in text

    def test_signal_list_value(self) -> None:
        doc = Document()
        _add_finance_section(doc, {
            "observation_signals": {
                "award_concentration": ["采购人A", "采购人B", "采购人C"],
            }
        })
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "采购人A" in text
        assert "采购人B" in text

    def test_signal_list_truncated_to_10(self) -> None:
        doc = Document()
        items = [f"item_{i}" for i in range(15)]
        _add_finance_section(doc, {
            "observation_signals": {
                "high_freq_cooccurrence": items,
            }
        })
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "item_0" in text
        assert "item_9" in text
        assert "item_14" not in text

    def test_signal_scalar_value(self) -> None:
        doc = Document()
        _add_finance_section(doc, {
            "observation_signals": {
                "explicit_rejection": "2次否决",
            }
        })
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "2次否决" in text

    def test_signal_none_value(self) -> None:
        doc = Document()
        _add_finance_section(doc, {
            "observation_signals": {
                "cancellation_link": None,
            }
        })
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "本期未观察到相关数据" in text

    def test_all_six_signals_rendered(self) -> None:
        doc = Document()
        _add_finance_section(doc, {
            "observation_signals": {
                "award_activity": 5,
                "award_concentration": "Top 3",
                "cancellation_link": 0,
                "explicit_rejection": 1,
                "info_conflict": "无冲突",
                "high_freq_cooccurrence": ["A", "B"],
            }
        })
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "中标活跃度" in text
        assert "公开中标集中度" in text
        assert "废标公告关联" in text
        assert "明确投标否决" in text
        assert "信息冲突观察" in text
        assert "高频共现提示" in text


class TestCoverWithFrequency:
    def test_cover_renders_frequency(self) -> None:
        filters = ParsedFilters(raw_query="test query", frequency="每天9:00")
        doc = Document()
        _add_cover(doc, filters, total=5)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "每天9:00" in text

    def test_cover_without_frequency(self) -> None:
        filters = ParsedFilters(raw_query="no freq")
        doc = Document()
        _add_cover(doc, filters, total=0)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "推送频率" not in text


class TestFilenameHelpers:
    def test_sanitize_removes_illegal_chars(self) -> None:
        result = _sanitize_filename('test<>:"/\\|?*file')
        for ch in '<>:"/\\|?*':
            assert ch not in result

    def test_sanitize_truncates_long_query(self) -> None:
        long_query = "A" * 120
        result = _sanitize_filename(long_query)
        assert len(result) == 80

    def test_sanitize_empty_returns_default(self) -> None:
        assert _sanitize_filename("") == "query"
        assert _sanitize_filename("   ") == "query"
        assert _sanitize_filename("...") == "query"

    def test_build_filename_format(self) -> None:
        dt = datetime(2026, 1, 15, 14, 30)
        filename = build_filename("测试查询", dt=dt)
        assert filename == "测试查询_202601151430.docx"

    def test_build_filename_sanitizes_query(self) -> None:
        dt = datetime(2026, 1, 1, 0, 0)
        filename = build_filename('test/file?name', dt=dt)
        assert "/" not in filename
        assert "?" not in filename
        assert filename.endswith("_202601010000.docx")


class TestGenerateReportSync:
    def test_generate_report_with_finance(self, tmp_path) -> None:
        filters = ParsedFilters(
            raw_query="测试报告",
            topic="服务器",
            region="上海",
            frequency="每天9:00",
        )
        items = [
            {
                "project_name": "测试项目",
                "publish_time": "2026-01-01T10:00:00",
                "source_url": "https://example.com/1",
                "core_content": "预算100万元",
                "attachment_url": "https://example.com/att.pdf",
                "budget_amount": 1000000,
                "source_platform": "ccgp",
            }
        ]
        finance = {"observation_signals": {"award_activity": 3}}

        with patch("app.report.docx_generator.settings") as mock_settings:
            mock_settings.REPORT_OUTPUT_DIR = str(tmp_path)
            filepath = _generate_report_sync(filters, items, finance_summary=finance)

        assert os.path.exists(filepath)
        assert filepath.endswith(".docx")

    def test_generate_report_empty_items(self, tmp_path) -> None:
        filters = ParsedFilters(raw_query="空测试")
        with patch("app.report.docx_generator.settings") as mock_settings:
            mock_settings.REPORT_OUTPUT_DIR = str(tmp_path)
            filepath = _generate_report_sync(filters, [])

        assert os.path.exists(filepath)
