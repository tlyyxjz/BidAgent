"""Extra tests for misc modules: docx_components / session_manager / logger / config.

Covers uncovered branches:
- docx_components: add_detail_table (empty, display_grade, publish_time parse),
  add_analysis (empty, top3, urgent deadlines), add_footer,
  add_anti_hallucination_section (passed / failed)
- session_manager: _cookie_expired (maxAge/expires/invalid), _domain_matches,
  cookie_summary with non-dict, delete non-existent, create_context without state
- logger: setup_logging, new_request_id, get_logger, RequestIdFilter
- config: validators (secret_key/admin_secret/free_limit/positive), proxies property
"""

from __future__ import annotations

import json
import time
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from app.core.session_manager import SessionManager
from app.utils.logger import (
    RequestIdFilter,
    get_logger,
    new_request_id,
    request_id_var,
    setup_logging,
)


# ============================================================
# Part 1: docx_components
# ============================================================

from app.report.docx_components import (
    add_analysis,
    add_anti_hallucination_section,
    add_detail_table,
    add_footer,
)
from app.report.docx_components import add_value_note


class TestAddDetailTable:
    """Cover add_detail_table branches."""

    def test_empty_items_shows_placeholder(self) -> None:
        doc = Document()
        add_detail_table(doc, [])
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "no data" in text or "暂无" in text

    def test_items_with_display_grade_high(self) -> None:
        doc = Document()
        items = [{
            "project_name": "测试项目",
            "publish_time": "2026-01-15T10:00:00",
            "source_url": "https://example.com/1",
            "core_content": "预算100万元",
            "attachment_url": "https://example.com/att.pdf",
            "display_grade": "high",
        }]
        add_detail_table(doc, items)
        # 表格应有数据行
        assert len(doc.tables) >= 1
        table_text = " ".join(
            cell.text for row in doc.tables[0].rows for cell in row.cells
        )
        assert "测试项目" in table_text
        assert "可信" in table_text

    def test_items_with_display_grade_review(self) -> None:
        doc = Document()
        items = [{
            "project_name": "待核项目",
            "publish_time": "2026-01-15",
            "source_url": "https://example.com/2",
            "core_content": "内容",
            "attachment_url": "",
            "display_grade": "review",
        }]
        add_detail_table(doc, items)
        table_text = " ".join(
            cell.text for row in doc.tables[0].rows for cell in row.cells
        )
        assert "待核" in table_text

    def test_items_with_display_grade_low(self) -> None:
        doc = Document()
        items = [{
            "project_name": "存疑项目",
            "publish_time": None,
            "source_url": "https://example.com/3",
            "core_content": "内容",
            "attachment_url": "",
            "display_grade": "low",
        }]
        add_detail_table(doc, items)
        table_text = " ".join(
            cell.text for row in doc.tables[0].rows for cell in row.cells
        )
        assert "存疑" in table_text
        assert "-" in table_text  # publish_time None -> "-"

    def test_items_without_display_grade(self) -> None:
        doc = Document()
        items = [{
            "project_name": "普通项目",
            "publish_time": "2026-01-15T10:00:00",
            "source_url": "https://example.com/4",
            "core_content": "普通内容",
            "attachment_url": "https://example.com/att.pdf",
        }]
        add_detail_table(doc, items)
        table_text = " ".join(
            cell.text for row in doc.tables[0].rows for cell in row.cells
        )
        assert "普通项目" in table_text
        # 无 grade 不应有徽标
        assert "可信" not in table_text
        assert "待核" not in table_text

    def test_publish_time_invalid_falls_back(self) -> None:
        doc = Document()
        items = [{
            "project_name": "项目",
            "publish_time": "not-a-date",
            "source_url": "https://example.com",
            "core_content": "内容",
            "attachment_url": "",
        }]
        add_detail_table(doc, items)
        table_text = " ".join(
            cell.text for row in doc.tables[0].rows for cell in row.cells
        )
        # 非法日期取前10字符
        assert "not-a-date"[:10] in table_text

    def test_publish_time_none_shows_dash(self) -> None:
        doc = Document()
        items = [{
            "project_name": "项目",
            "publish_time": None,
            "source_url": "https://example.com",
            "core_content": "内容",
            "attachment_url": "",
        }]
        add_detail_table(doc, items)
        table_text = " ".join(
            cell.text for row in doc.tables[0].rows for cell in row.cells
        )
        assert "-" in table_text


class TestAddAnalysis:
    """Cover add_analysis branches."""

    def test_empty_items_shows_placeholder(self) -> None:
        doc = Document()
        add_analysis(doc, [])
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "no data" in text or "暂无" in text

    def test_items_with_budget_sorting(self) -> None:
        doc = Document()
        items = [
            {"project_name": "A", "budget_amount": 100000, "tender_org": "甲方A"},
            {"project_name": "B", "budget_amount": 500000, "tender_org": "甲方B"},
            {"project_name": "C", "budget_amount": 300000, "tender_org": "甲方C"},
        ]
        add_analysis(doc, items)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "B" in text  # 最高预算排第一

    def test_items_with_deadline(self) -> None:
        doc = Document()
        items = [
            {
                "project_name": "紧急项目",
                "budget_amount": 100000,
                "deadline": "2026-12-31T23:59",
            },
        ]
        add_analysis(doc, items)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "即将截止" in text or "截止" in text

    def test_items_with_invalid_deadline(self) -> None:
        doc = Document()
        items = [
            {
                "project_name": "项目",
                "budget_amount": 100000,
                "deadline": "invalid-date",
            },
        ]
        add_analysis(doc, items)
        # 不应崩溃
        text = "\n".join(p.text for p in doc.paragraphs)
        assert len(text) > 0

    def test_items_without_deadline(self) -> None:
        doc = Document()
        items = [
            {"project_name": "项目", "budget_amount": 100000},
        ]
        add_analysis(doc, items)
        text = "\n".join(p.text for p in doc.paragraphs)
        # 无 deadline 不显示截止提醒
        assert "即将截止" not in text

    def test_grade_description_present(self) -> None:
        doc = Document()
        items = [{"project_name": "项目", "budget_amount": 100000}]
        add_analysis(doc, items)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "可信度" in text or "high" in text


class TestAddFooter:
    """Cover add_footer."""

    def test_footer_contains_text(self) -> None:
        doc = Document()
        add_footer(doc)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "报告结束" in text or "标小智" in text or "标" in text

    def test_footer_has_date(self) -> None:
        doc = Document()
        add_footer(doc)
        text = "\n".join(p.text for p in doc.paragraphs)
        # 应包含日期格式
        assert "20" in text  # 年份前缀


class TestAddAntiHallucinationSection:
    """Cover add_anti_hallucination_section passed/failed branches."""

    def test_all_passed(self) -> None:
        doc = Document()
        items = [{
            "project_name": "项目",
            "core_content": "预算100万元",
            "source_url": "https://example.com/1",
        }]
        source_texts = {"https://example.com/1": "预算100万元"}
        add_anti_hallucination_section(doc, items, source_texts=source_texts)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "总项目数" in text

    def test_empty_items(self) -> None:
        doc = Document()
        add_anti_hallucination_section(doc, [])
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "总项目数" in text
        assert "0" in text

    def test_with_failed_items(self) -> None:
        doc = Document()
        items = [{
            "project_name": "幻觉项目",
            "core_content": "预算999万元 不存在的编号XYZ",
            "source_url": "https://example.com/2",
        }]
        # source_text 不包含这些事实，应触发幻觉
        source_texts = {"https://example.com/2": "完全不同的内容"}
        add_anti_hallucination_section(doc, items, source_texts=source_texts)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "总项目数" in text


# ============================================================
# Part 2: session_manager
# ============================================================


def _write_state(path: Path, cookies=None, origins=None) -> None:
    path.write_text(
        json.dumps({"cookies": cookies or [], "origins": origins or []}),
        encoding="utf-8",
    )


class TestCookieExpired:
    """Cover SessionManager._cookie_expired branches."""

    def test_max_age_zero_expired(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"maxAge": 0}, now) is True

    def test_max_age_negative_expired(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"maxAge": -1}, now) is True

    def test_max_age_invalid_returns_true(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"maxAge": "abc"}, now) is True

    def test_max_age_positive_not_expired(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"maxAge": 3600}, now) is False

    def test_max_age_hyphen_key(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"max-age": 0}, now) is True
        assert SessionManager._cookie_expired({"max-age": 3600}, now) is False

    def test_expires_future_not_expired(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"expires": now + 100}, now) is False

    def test_expires_past_expired(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"expires": now - 100}, now) is True

    def test_expires_session_cookie_not_expired(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"expires": -1}, now) is False
        assert SessionManager._cookie_expired({"expires": 0}, now) is False

    def test_expires_invalid_returns_true(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({"expires": "abc"}, now) is True

    def test_expires_none_not_expired(self) -> None:
        now = time.time()
        assert SessionManager._cookie_expired({}, now) is False


class TestDomainMatches:
    """Cover SessionManager._domain_matches."""

    def test_exact_match(self) -> None:
        assert SessionManager._domain_matches("example.com", "example.com") is True

    def test_subdomain_match(self) -> None:
        assert SessionManager._domain_matches("www.example.com", "example.com") is True

    def test_dot_prefix_match(self) -> None:
        assert SessionManager._domain_matches(".example.com", "example.com") is True

    def test_no_match(self) -> None:
        assert SessionManager._domain_matches("evilexample.com", "example.com") is False

    def test_case_insensitive(self) -> None:
        assert SessionManager._domain_matches("EXAMPLE.COM", "example.com") is True
        assert SessionManager._domain_matches("WWW.Example.com", "example.com") is True


class TestSessionManagerExtra:
    """Cover extra session_manager branches."""

    @pytest.mark.asyncio
    async def test_cookie_summary_empty_state(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        _write_state(path, [])
        summary = await SessionManager("test", path).cookie_summary()
        assert summary == []

    @pytest.mark.asyncio
    async def test_cookie_summary_skips_non_dict(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        path.write_text(
            json.dumps({"cookies": ["not-a-dict", {"name": "ok"}], "origins": []}),
            encoding="utf-8",
        )
        summary = await SessionManager("test", path).cookie_summary()
        assert len(summary) == 1
        assert summary[0]["name"] == "ok"

    @pytest.mark.asyncio
    async def test_is_valid_empty_state(self, tmp_path: Path) -> None:
        path = tmp_path / "missing.json"
        assert await SessionManager("test", path).is_valid() is False

    @pytest.mark.asyncio
    async def test_is_valid_no_required_cookies(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        _write_state(path, [{"name": "sid", "domain": ".example.com", "expires": -1}])
        assert await SessionManager("test", path).is_valid() is True

    @pytest.mark.asyncio
    async def test_is_valid_skips_non_dict_cookie(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        path.write_text(
            json.dumps({
                "cookies": ["bad", {"name": "good", "domain": ".x.com", "expires": -1}],
                "origins": [],
            }),
            encoding="utf-8",
        )
        assert await SessionManager("test", path).is_valid() is True

    @pytest.mark.asyncio
    async def test_is_valid_skips_empty_name(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        _write_state(path, [{"name": "", "domain": ".x.com", "expires": -1}])
        assert await SessionManager("test", path).is_valid() is False

    @pytest.mark.asyncio
    async def test_create_context_without_state(self, tmp_path: Path) -> None:
        from unittest.mock import AsyncMock
        path = tmp_path / "missing.json"
        browser = AsyncMock()
        await SessionManager("test", path).create_context(browser, locale="en")
        browser.new_context.assert_awaited_once_with(locale="en")

    @pytest.mark.asyncio
    async def test_delete_nonexistent_returns_false(self, tmp_path: Path) -> None:
        path = tmp_path / "missing.json"
        assert await SessionManager("test", path).delete() is False

    @pytest.mark.asyncio
    async def test_delete_existing_returns_true(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        _write_state(path)
        assert await SessionManager("test", path).delete() is True
        assert not path.exists()

    @pytest.mark.asyncio
    async def test_load_state_invalid_cookies_type(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        path.write_text(
            json.dumps({"cookies": "not-a-list", "origins": []}),
            encoding="utf-8",
        )
        state = await SessionManager("test", path).load_state()
        assert state is None

    @pytest.mark.asyncio
    async def test_load_state_invalid_origins_type(self, tmp_path: Path) -> None:
        path = tmp_path / "state.json"
        path.write_text(
            json.dumps({"cookies": [], "origins": "not-a-list"}),
            encoding="utf-8",
        )
        state = await SessionManager("test", path).load_state()
        assert state is None

    @pytest.mark.asyncio
    async def test_save_invalid_state_type_raises(self, tmp_path: Path) -> None:
        from unittest.mock import AsyncMock
        path = tmp_path / "state.json"
        ctx = AsyncMock()
        ctx.storage_state.return_value = "not-a-dict"
        with pytest.raises(ValueError):
            await SessionManager("test", path).save(ctx)

    @pytest.mark.asyncio
    async def test_save_invalid_cookies_type_raises(self, tmp_path: Path) -> None:
        from unittest.mock import AsyncMock
        path = tmp_path / "state.json"
        ctx = AsyncMock()
        ctx.storage_state.return_value = {"cookies": "bad", "origins": []}
        with pytest.raises(ValueError):
            await SessionManager("test", path).save(ctx)

    @pytest.mark.asyncio
    async def test_save_invalid_origins_type_raises(self, tmp_path: Path) -> None:
        from unittest.mock import AsyncMock
        path = tmp_path / "state.json"
        ctx = AsyncMock()
        ctx.storage_state.return_value = {"cookies": [], "origins": "bad"}
        with pytest.raises(ValueError):
            await SessionManager("test", path).save(ctx)

    def test_platform_name_too_long_rejected(self) -> None:
        with pytest.raises(ValueError):
            SessionManager("x" * 100)

    def test_platform_name_with_special_char_rejected(self) -> None:
        with pytest.raises(ValueError):
            SessionManager("bad/platform")


# ============================================================
# Part 3: logger
# ============================================================


class TestSetupLogging:
    """Cover setup_logging function."""

    def test_setup_logging_default_level(self) -> None:
        # Should not raise
        setup_logging()
        logger = get_logger("test_setup")
        assert logger is not None

    def test_setup_logging_custom_level(self) -> None:
        setup_logging(level="DEBUG")
        logger = get_logger("test_debug")
        assert logger is not None

    def test_setup_logging_idempotent(self) -> None:
        setup_logging()
        setup_logging()
        logger = get_logger("test_idem")
        assert logger is not None


class TestNewRequestId:
    """Cover new_request_id function."""

    def test_returns_non_empty_string(self) -> None:
        rid = new_request_id()
        assert isinstance(rid, str)
        assert len(rid) == 12

    def test_sets_context_var(self) -> None:
        rid = new_request_id()
        assert request_id_var.get() == rid

    def test_generates_unique_ids(self) -> None:
        ids = {new_request_id() for _ in range(100)}
        assert len(ids) == 100


class TestGetLogger:
    """Cover get_logger function."""

    def test_returns_bound_logger(self) -> None:
        logger = get_logger("my_module")
        assert logger is not None
        # Should be callable for logging
        logger.info("test message")

    def test_default_name(self) -> None:
        logger = get_logger()
        assert logger is not None


class TestRequestIdFilter:
    """Cover RequestIdFilter."""

    def test_filter_sets_request_id(self) -> None:
        new_request_id()
        f = RequestIdFilter()
        import logging
        record = logging.LogRecord(
            name="test", level=logging.INFO, pathname="", lineno=0,
            msg="test", args=(), exc_info=None,
        )
        assert f.filter(record) is True
        assert hasattr(record, "request_id")
        assert record.request_id == request_id_var.get()


# ============================================================
# Part 4: config
# ============================================================


class TestConfigValidators:
    """Cover config validators."""

    def test_invalid_secret_key_non_hex_raises(self) -> None:
        from app.config import Settings
        from pydantic import ValidationError
        with patch.dict("os.environ", {
            "SECRET_KEY": "not-hex-string!",
            "ADMIN_SECRET": "test-admin-12345",
        }, clear=False):
            with pytest.raises((ValidationError, ValueError)):
                Settings()

    def test_invalid_secret_key_wrong_length_raises(self) -> None:
        from app.config import Settings
        from pydantic import ValidationError
        # 16 chars hex = 8 bytes, not 32 bytes
        with patch.dict("os.environ", {
            "SECRET_KEY": "abcdef0123456789",
            "ADMIN_SECRET": "test-admin-12345",
        }, clear=False):
            with pytest.raises((ValidationError, ValueError)):
                Settings()

    def test_empty_admin_secret_raises(self) -> None:
        from app.config import Settings
        from pydantic import ValidationError
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "   ",
        }, clear=False):
            with pytest.raises((ValidationError, ValueError)):
                Settings()

    def test_short_admin_secret_raises(self) -> None:
        from app.config import Settings
        from pydantic import ValidationError
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "short",
        }, clear=False):
            with pytest.raises((ValidationError, ValueError)):
                Settings()

    def test_zero_free_limit_raises(self) -> None:
        from app.config import Settings
        from pydantic import ValidationError
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "FREE_TIER_DAILY_LIMIT": "0",
        }, clear=False):
            with pytest.raises((ValidationError, ValueError)):
                Settings()

    def test_zero_smtp_port_raises(self) -> None:
        from app.config import Settings
        from pydantic import ValidationError
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "SMTP_PORT": "0",
        }, clear=False):
            with pytest.raises((ValidationError, ValueError)):
                Settings()

    def test_negative_browser_pool_size_raises(self) -> None:
        from app.config import Settings
        from pydantic import ValidationError
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "BROWSER_POOL_SIZE": "-1",
        }, clear=False):
            with pytest.raises((ValidationError, ValueError)):
                Settings()


class TestConfigProperties:
    """Cover config property accessors."""

    def test_proxies_empty(self) -> None:
        from app.config import Settings
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "PROXY_LIST": "",
        }, clear=False):
            s = Settings()
            assert s.proxies == []

    def test_proxies_with_values(self) -> None:
        from app.config import Settings
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "PROXY_LIST": "http://p1.com, http://p2.com ,",
        }, clear=False):
            s = Settings()
            assert s.proxies == ["http://p1.com", "http://p2.com"]

    def test_proxies_whitespace_only(self) -> None:
        from app.config import Settings
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "PROXY_LIST": "   ,  ,  ",
        }, clear=False):
            s = Settings()
            assert s.proxies == []

    def test_cors_origin_list_explicit(self) -> None:
        from app.config import Settings
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "CORS_ORIGINS": "https://a.com, https://b.com ",
        }, clear=False):
            s = Settings()
            assert s.cors_origin_list == ["https://a.com", "https://b.com"]

    def test_cors_origin_list_empty(self) -> None:
        from app.config import Settings
        with patch.dict("os.environ", {
            "SECRET_KEY": "a" * 64,
            "ADMIN_SECRET": "test-admin-12345",
            "CORS_ORIGINS": "",
        }, clear=False):
            s = Settings()
            assert s.cors_origin_list == []

    def test_secret_key_lowercased(self) -> None:
        from app.config import Settings
        with patch.dict("os.environ", {
            "SECRET_KEY": "A" * 64,
            "ADMIN_SECRET": "test-admin-12345",
        }, clear=False):
            s = Settings()
            assert s.SECRET_KEY == "a" * 64

    def test_valid_settings_load_successfully(self) -> None:
        from app.config import Settings
        with patch.dict("os.environ", {
            "SECRET_KEY": "b" * 64,
            "ADMIN_SECRET": "valid-admin-12345",
            "FREE_TIER_DAILY_LIMIT": "10",
            "SMTP_PORT": "587",
            "BROWSER_POOL_SIZE": "3",
        }, clear=False):
            s = Settings()
            assert s.SECRET_KEY == "b" * 64
            assert s.FREE_TIER_DAILY_LIMIT == 10
            assert s.SMTP_PORT == 587
            assert s.BROWSER_POOL_SIZE == 3


def test_add_value_note_cost_narrative() -> None:
    """成本叙事：价值脚注包含实测口径数字（0.85 分钱 / 24 秒）。"""
    doc = Document()
    add_value_note(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "0.85 分钱" in text
    assert "24 秒" in text
    assert "1-2 人时" in text
    assert "自动生成" in text
